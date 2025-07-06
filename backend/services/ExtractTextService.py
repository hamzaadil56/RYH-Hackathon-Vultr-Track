import os
import fitz
from docx import Document
from fastapi import UploadFile
import re
from io import BytesIO


class ExtractTextService:
    def extract_text_from_file(self,file: UploadFile) -> str:

        file_extension = os.path.splitext(file.filename)[1].lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: .pdf, .docx")
    
    def extract_text_from_pdf(self,file: UploadFile) -> str:
        extracted_text = ""
        
        try:
            # Read file content as bytes
            file_bytes = file.file.read()
            
            # Reset file pointer if needed for future reads
            file.file.seek(0)
            
            pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
            
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                page_text = page.get_text()
                extracted_text += page_text + "\n"
            
            pdf_document.close()
            return extracted_text.strip()
                
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")

    def extract_text_from_docx(self,file: UploadFile) -> str:
        extracted_text = ""
        try:
            # Read file content as bytes and create BytesIO stream
            file_bytes = file.file.read()
            
            # Reset file pointer if needed for future reads
            file.file.seek(0)
            
            doc = Document(BytesIO(file_bytes))
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                extracted_text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text += cell.text + " "
                    extracted_text += "\n"
            
            return extracted_text.strip()
            
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    def clean_extracted_text(self,text: str) -> str:
    
        lines = [line.strip() for line in text.split('\n')]
        lines = [line for line in lines if line]  # Remove empty lines
        
        # Join lines with single newlines
        cleaned_text = '\n'.join(lines)
        
        # Remove multiple consecutive spaces
        
        cleaned_text = re.sub(r' +', ' ', cleaned_text)
        
        return cleaned_text

    def extract_and_clean_text(self,file: UploadFile) -> str:
        raw_text = self.extract_text_from_file(file)
        return self.clean_extracted_text(raw_text)